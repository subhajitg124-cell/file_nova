import os
import pytest
from pathlib import Path
from PIL import Image
from app.utils.mime import detect_mime_type
from app.processors.image import ImageProcessor
from app.processors.pdf import PDFProcessor
from app.config import settings

@pytest.fixture
def temp_image_file(tmp_path):
    """Fixture to create a tiny temporary PNG image file for tests."""
    img_path = tmp_path / "test_input.png"
    img = Image.new("RGBA", (100, 100), color=(255, 0, 0, 128))
    img.save(img_path, "PNG")
    return img_path

@pytest.fixture
def temp_text_file(tmp_path):
    """Fixture to create a temporary markdown file."""
    txt_path = tmp_path / "test_doc.md"
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write("# Hello World\nThis is a test paragraph.")
    return txt_path


def test_magic_bytes_detection(temp_image_file, temp_text_file):
    """Verifies that our custom pure-python magic detector operates correctly."""
    # Test PNG
    mime_png = detect_mime_type(temp_image_file)
    assert mime_png == "image/png"
    
    # Test MD
    mime_md = detect_mime_type(temp_text_file)
    assert mime_md == "text/plain"  # Decoded plain text fallback


def test_image_compress_processor(temp_image_file):
    """Tests that ImageProcessor runs compression operations and writes output."""
    settings.output_dir.mkdir(parents=True, exist_ok=True)
    
    proc = ImageProcessor()
    options = {
        "operation": "compress",
        "quality": 70,
        "resize_pct": 0.5,
        "target_format": "jpeg"
    }
    
    out_path_str = proc.process([str(temp_image_file)], options)
    out_path = Path(out_path_str)
    
    assert out_path.exists()
    assert out_path.suffix.lower() == ".jpg"
    
    # Verify dimensions are scaled down (100 * 0.5 = 50)
    with Image.open(out_path) as out_img:
        assert out_img.width == 50
        assert out_img.height == 50
        assert out_img.mode == "RGB"  # Discarded transparency channel for JPEG conversion


def test_image_enhance_processor(temp_image_file):
    """Tests that ImageProcessor runs enhancement matrix filters."""
    proc = ImageProcessor()
    options = {
        "operation": "enhance",
        "brightness": 1.5,
        "contrast": 1.2,
        "sharpness": 1.1,
        "denoise": True,
        "target_format": "png"
    }
    
    out_path_str = proc.process([str(temp_image_file)], options)
    out_path = Path(out_path_str)
    
    assert out_path.exists()
    assert out_path.suffix.lower() == ".png"


def test_docx_merge_processor(tmp_path):
    """Tests that DocumentProcessor can merge multiple DOCX files successfully."""
    from docx import Document
    from app.processors.docx_pdf import DocumentProcessor
    
    settings.output_dir.mkdir(parents=True, exist_ok=True)
    
    # Create two dummy docx files
    doc1_path = tmp_path / "doc1.docx"
    doc2_path = tmp_path / "doc2.docx"
    
    doc1 = Document()
    doc1.add_heading("Doc 1 Heading", level=1)
    doc1.add_paragraph("This is the first document text.")
    doc1.save(doc1_path)
    
    doc2 = Document()
    doc2.add_heading("Doc 2 Heading", level=1)
    doc2.add_paragraph("This is the second document text.")
    doc2.save(doc2_path)
    
    # Run processor
    proc = DocumentProcessor()
    options = {
        "operation": "docx_merge"
    }
    
    out_path_str = proc.process([str(doc1_path), str(doc2_path)], options)
    out_path = Path(out_path_str)
    
    assert out_path.exists()
    assert out_path.name.startswith("merged_")
    assert out_path.suffix.lower() == ".docx"
    
    # Load and verify merged content
    merged_doc = Document(out_path)
    all_text = [p.text for p in merged_doc.paragraphs]
    
    assert "Doc 1 Heading" in all_text
    assert "This is the first document text." in all_text
    assert "Doc 2 Heading" in all_text
    assert "This is the second document text." in all_text

